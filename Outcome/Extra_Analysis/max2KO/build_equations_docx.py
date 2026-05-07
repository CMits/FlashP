#!/usr/bin/env python3
"""
Build a Word document for Figure 2A:

  - Two side-by-side tables, each 2 columns wide:
      Table 1 (WT)        : Equation | WT calculation
      Table 2 (MAX2 KO)   : Equation | MAX2-KO calculation
  - All equations are real Word equations (Office Math / OMML), so subscripts,
    superscripts, and fractions render properly — not inline ASCII.
  - Below the tables, a compact algebraic vs ODE steady-state summary.

Outputs:
    docs/MAX2_equations_and_values.docx
"""
from __future__ import annotations

import datetime
import json
import math
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Union

from lxml import etree

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT

ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent
NET_DIR = REPO / "Arabidopsis" / "Shoot_Branching_network" / "network"
VAL_DIR = REPO / "Arabidopsis" / "Shoot_Branching_network" / "validation"

ALG_JSON = NET_DIR / "algebraic_equations.json"
ALG_DUMP = VAL_DIR / "steady_state_dump.json"
ODE_DUMP = VAL_DIR / "ode_steady_state_dump.json"
WT_KO_JSON = ROOT / "data" / "max2_ko_steady_state.json"

OUT_DOCX = ROOT / "docs" / "MAX2_equations_and_values.docx"

SUBGRAPH_NODES = [
    "Strigolactone", "D14", "MAX2", "SMXL678", "BES1", "SPL9", "BRC1",
    "HB21", "NCED3", "ABA", "PIN1", "PIN3", "Shoot_Branching",
]
TEST_ID = "T005"

ACTIVATOR_FLOOR = 0.01
INHIB_FLOOR = 0.1
INHIB_CAP = 10.0


def _n(x: float) -> str:
    """Format a constant compactly: 10.0 -> '10', 0.1 -> '0.1'."""
    if x == int(x):
        return str(int(x))
    s = f"{x:g}"
    return s


# ---------------------------------------------------------------------------
# Equation tree (mini-AST) — leaf = Text, structural = Sub / Sup / Frac / Seq
# ---------------------------------------------------------------------------

@dataclass
class Text:
    s: str

@dataclass
class Sub:
    base: "Node"
    sub: "Node"

@dataclass
class Sup:
    base: "Node"
    sup: "Node"

@dataclass
class Frac:
    num: "Node"
    den: "Node"

@dataclass
class Seq:
    items: list = field(default_factory=list)


Node = Union[Text, Sub, Sup, Frac, Seq, str]


# ---------------------------------------------------------------------------
# OMML serialisation
# ---------------------------------------------------------------------------

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
M = f"{{{OMML_NS}}}"
NSMAP = {"m": OMML_NS}
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _omml_run(text: str) -> etree._Element:
    r = etree.Element(M + "r")
    t = etree.SubElement(r, M + "t")
    t.text = text
    t.set(XML_SPACE, "preserve")
    return r


def _to_omml(node: Node) -> list:
    if isinstance(node, str):
        return [_omml_run(node)]
    if isinstance(node, Text):
        return [_omml_run(node.s)]
    if isinstance(node, Seq):
        out = []
        for it in node.items:
            out.extend(_to_omml(it))
        return out
    if isinstance(node, Sub):
        s = etree.Element(M + "sSub")
        e = etree.SubElement(s, M + "e")
        for el in _to_omml(node.base):
            e.append(el)
        sub = etree.SubElement(s, M + "sub")
        for el in _to_omml(node.sub):
            sub.append(el)
        return [s]
    if isinstance(node, Sup):
        s = etree.Element(M + "sSup")
        e = etree.SubElement(s, M + "e")
        for el in _to_omml(node.base):
            e.append(el)
        sup = etree.SubElement(s, M + "sup")
        for el in _to_omml(node.sup):
            sup.append(el)
        return [s]
    if isinstance(node, Frac):
        f = etree.Element(M + "f")
        num = etree.SubElement(f, M + "num")
        for el in _to_omml(node.num):
            num.append(el)
        den = etree.SubElement(f, M + "den")
        for el in _to_omml(node.den):
            den.append(el)
        return [f]
    raise ValueError(f"Unknown node type: {type(node)}")


def _make_omath(tree: Node) -> etree._Element:
    o = etree.Element(M + "oMath", nsmap=NSMAP)
    for el in _to_omml(tree):
        o.append(el)
    return o


def write_equation_cell(cell, tree: Node) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para._p.append(_make_omath(tree))


# ---------------------------------------------------------------------------
# Helpers for compact value formatting
# ---------------------------------------------------------------------------

def short(x: float) -> str:
    if abs(x) < 1e-6:
        return "0"
    if abs(x - round(x)) < 1e-4 and abs(x) < 100:
        return f"{int(round(x))}"
    if abs(x) < 1e-3:
        return f"{x:.1e}"
    if abs(x) >= 100:
        return f"{x:.1f}"
    return f"{x:.3f}".rstrip("0").rstrip(".")


def fmt(x: float) -> str:
    if abs(x) < 1e-3:
        return f"{x:.3e}"
    if abs(x - round(x)) < 1e-4 and abs(x) < 1000:
        return f"{int(round(x))}"
    if abs(x) >= 100:
        return f"{x:.1f}"
    return f"{x:.3f}"


# ---------------------------------------------------------------------------
# Equation builders → AST
# ---------------------------------------------------------------------------

def sym_eq(eq: dict) -> Seq:
    """Symbolic equation as an AST."""
    node = eq["node"]
    if eq.get("is_source"):
        return Seq([
            Text(f"{node} = "),
            Sub(Text("g"), Text(node)),
            Text(" + e"),
        ])
    items: list = [Text(f"{node} = ")]
    acts, inhs = eq["activators"], eq["inhibitors"]
    if acts:
        if len(acts) == 1:
            items.append(Text(f"max({acts[0]}, {ACTIVATOR_FLOOR})"))
        else:
            inner = " · ".join(f"max({a}, {ACTIVATOR_FLOOR})" for a in acts)
            items.append(Sup(Text(f"({inner})"), Text(f"1/{len(acts)}")))
        items.append(Text(" · "))
    if inhs:
        denom = " · ".join(inhs) if len(inhs) > 1 else inhs[0]
        items.append(Text("min("))
        items.append(Frac(Text("1"), Text(f"max({denom}, {INHIB_FLOOR})")))
        items.append(Text(f", {_n(INHIB_CAP)}) · "))
    items.append(Sub(Text("g"), Text(node)))
    items.append(Text(" + e"))
    return Seq(items)


def calc_eq(eq: dict, values: dict, g_for_node: float) -> Seq:
    """Substituted calculation as an AST.

    Mirrors the FLASH-P algebraic step:
        f = Activation · Inhibition · g + e
        Activation = ( ∏ max(x, ε_floor) )^(1/n)        (n>=2; otherwise the bare value)
        Inhibition = min( 1 / max( ∏ x , ε ) , K )
    """
    node = eq["node"]
    if eq.get("is_source"):
        return Seq([
            Text(f"{node} = "),
            Sub(Text("g"), Text(node)),
            Text(f" + 0 = {short(g_for_node)} + 0 = {short(g_for_node)}"),
        ])
    items: list = [Text(f"{node} = ")]
    acts, inhs = eq["activators"], eq["inhibitors"]

    # Activation factor
    if acts:
        a_vals = [max(values[a], ACTIVATOR_FLOOR) for a in acts]
        if len(acts) == 1:
            a_factor = a_vals[0]
            items.append(Text(short(a_factor)))
        else:
            prod = 1.0
            for v in a_vals:
                prod *= v
            a_factor = prod ** (1.0 / len(acts))
            inner = " · ".join(short(v) for v in a_vals)
            items.append(Sup(Text(f"({inner})"), Text(f"1/{len(acts)}")))
            items.append(Text(f" = {short(a_factor)}"))
        items.append(Text(" · "))
    else:
        a_factor = 1.0

    # Inhibition factor
    if inhs:
        i_vals = [values[h] for h in inhs]
        prod = 1.0
        for v in i_vals:
            prod *= v
        denom = max(prod, INHIB_FLOOR)
        i_factor = min(1.0 / denom, INHIB_CAP)
        inner = " · ".join(short(v) for v in i_vals) if len(inhs) > 1 else short(i_vals[0])
        items.append(Text("min("))
        items.append(Frac(Text("1"), Text(f"max({inner}, {INHIB_FLOOR})")))
        items.append(Text(f", {_n(INHIB_CAP)}) = {short(i_factor)} · "))
    else:
        i_factor = 1.0

    final = a_factor * i_factor * g_for_node + 0.0
    items.append(Text(f"{short(g_for_node)} + 0 = {short(final)}"))
    return Seq(items)


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def load_dump_block(path: Path, test_id: str) -> dict:
    for block in json.loads(path.read_text(encoding="utf-8")):
        if block.get("test_id") == test_id:
            return block
    raise RuntimeError(f"{test_id} not in {path}")


def build_calc_table(doc, eq_by_node: dict, values: dict,
                     g_overrides: dict, title: str, calc_header: str) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    table.allow_autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = "Equation"
    hdr[1].text = calc_header
    for c in hdr:
        for run in c.paragraphs[0].runs:
            run.bold = True
    # Landscape A4 with 1.5 cm margins ≈ 26.7 cm printable width.
    widths = [Cm(10.0), Cm(16.5)]
    for i, c in enumerate(hdr):
        c.width = widths[i]

    for node in SUBGRAPH_NODES:
        eq = eq_by_node.get(node)
        if eq is None:
            continue
        g_node = g_overrides.get(node, 1.0)
        sym_tree = sym_eq(eq)
        calc_tree = calc_eq(eq, values, g_node)
        row = table.add_row().cells
        write_equation_cell(row[0], sym_tree)
        write_equation_cell(row[1], calc_tree)
        for i, cell in enumerate(row):
            cell.width = widths[i]


def main() -> int:
    eqs_full = json.loads(ALG_JSON.read_text(encoding="utf-8"))
    params = eqs_full["parameters"]
    eq_by_node = {e["node"]: e for e in eqs_full["equations"]}

    alg = load_dump_block(ALG_DUMP, TEST_ID)["steady_state_values"]
    ode = load_dump_block(ODE_DUMP, TEST_ID)["steady_state_values"]
    run = json.loads(WT_KO_JSON.read_text(encoding="utf-8"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Landscape A4 with narrow margins, so the equations + calculation column
    # have plenty of horizontal room and nothing gets cut off.
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_w
    section.page_height = new_h
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    doc.add_heading("Figure 2A — MAX2 KO propagation: equations and node values", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Network: Arabidopsis shoot branching (38 nodes / 75 edges, "
        "iteration 4, refinement 1). Subgraph: 13 nodes on the directed "
        "cone from MAX2 to Shoot_Branching. Test "
    )
    p.add_run(TEST_ID).bold = True
    p.add_run(" = MAX2 single knockout. Predicted increased; expected increased; CORRECT.")

    doc.add_heading("Global rules", level=2)
    rules = [
        "Activation (geometric mean):  A = ( ∏ max(x_a, ε_floor) )^(1/n)",
        "Inhibition (bounded inverse):  I = min( 1 / max( ∏ x_h , ε ) , K )",
        "Update:  f = A · I · g + e        (source: f = g + e)",
        "Damped Jacobi step:  x^(t) = (1−λ) · x^(t−1) + λ · f(x^(t−1))",
        "Convergence:  max_i |x_i^(t) − x_i^(t−1)| < δ",
    ]
    for r in rules:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph(
        f"Parameters: λ = {_n(params['damping'])}, δ = {params['convergence_tolerance']}, "
        f"K = {_n(params['K'])}, ε = {_n(params['epsilon'])}, "
        f"ε_floor = {_n(params['activator_floor'])}. "
        f"This run converged at iteration {run['converged_at_iteration']}; "
        f"final Shoot_Branching = {fmt(run['ko_run']['Shoot_Branching'])} "
        f"(WT baseline = 1.000). Cross-check vs steady_state_dump.json: "
        f"max abs diff {run['max_abs_diff_vs_dump']['value']:.2e} on "
        f"{run['max_abs_diff_vs_dump']['node']}."
    )

    # Table 1: WT
    wt_state = {n: 1.0 for n in eq_by_node}
    build_calc_table(
        doc, eq_by_node,
        values=wt_state,
        g_overrides={},  # all g = 1
        title="Table 1 — Wild type (WT)",
        calc_header="WT calculation",
    )

    # Table 2: MAX2 KO
    ko_state = run["ko_run"]
    build_calc_table(
        doc, eq_by_node,
        values=ko_state,
        g_overrides={"MAX2": 0.0},
        title="Table 2 — MAX2 knockout",
        calc_header="MAX2-KO calculation",
    )

    # Compact summary
    doc.add_heading("Steady-state summary (algebraic vs ODE)", level=2)
    summary = doc.add_table(rows=1, cols=6)
    summary.style = "Light Grid Accent 1"
    sh = summary.rows[0].cells
    for i, h in enumerate(["Node", "WT (alg)", "KO (alg)", "log2 FC (alg)",
                           "WT (ODE)", "KO (ODE)"]):
        sh[i].text = h
        sh[i].paragraphs[0].runs[0].bold = True
    for node in SUBGRAPH_NODES:
        wt_a = alg["WT"].get(node, 1.0)
        ko_a = alg["perturbed"].get(node, 1.0)
        wt_o = ode["WT"].get(node, 1.0)
        ko_o = ode["perturbed"].get(node, 1.0)
        l2 = math.log2(max(ko_a, 1e-6) / max(wt_a, 1e-6))
        row = summary.add_row().cells
        row[0].text = node
        row[1].text = fmt(wt_a)
        row[2].text = fmt(ko_a)
        row[3].text = f"{l2:+.2f}"
        row[4].text = fmt(wt_o)
        row[5].text = fmt(ko_o)
        if node == "Shoot_Branching":
            for c in row:
                for para in c.paragraphs:
                    for r in para.runs:
                        r.bold = True

    doc.add_heading("Files / provenance", level=2)
    for line in [
        f"Equations: {ALG_JSON.relative_to(REPO).as_posix()}",
        f"Algebraic dump: {ALG_DUMP.relative_to(REPO).as_posix()}",
        f"ODE dump: {ODE_DUMP.relative_to(REPO).as_posix()}",
        "Iteration trace: PropagationFigure/data/max2_ko_iteration_trace.csv",
        "Subgraph (Cytoscape): PropagationFigure/cytoscape/subgraph.graphml",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    target = OUT_DOCX
    try:
        doc.save(target)
    except PermissionError:
        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        target = OUT_DOCX.with_name(f"{OUT_DOCX.stem}_{stamp}{OUT_DOCX.suffix}")
        doc.save(target)
        print(
            f"WARNING: {OUT_DOCX.name} was locked (likely open in Word). "
            f"Wrote {target.name} instead."
        )
    print(f"Wrote {target}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
